"""
Document generation with random charts as images using python-docx-template.

This module demonstrates how to generate random pie charts and bar charts,
save them as PNG files, and embed them in a Word document.
"""
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import random
import os

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import logging
from matplotlib.figure import Figure
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def generate_pie_chart(output_path: str, title: str = 'Sample Pie Chart') -> None:
    """Generate a random pie chart and save it as a PNG file.
    
    Args:
        output_path: Path where to save the generated chart.
        title: Title of the chart.
    """
    # Generate random data
    sizes = [random.randint(5, 20) for _ in range(5)]
    labels = [f'Category {i+1}' for i in range(len(sizes))]
    colors = plt.cm.Pastel1(range(len(sizes)))
    
    # Create figure
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    ax.set_title(title)
    
    # Save the figure
    plt.tight_layout()
    fig.savefig(output_path, dpi=100, bbox_inches='tight')
    plt.close()

def generate_bar_chart(output_path: str, title: str = 'Sample Bar Chart') -> None:
    """Generate a random bar chart and save it as a PNG file.
    
    Args:
        output_path: Path where to save the generated chart.
        title: Title of the chart.
    """
    # Generate random data
    categories = [f'Q{i+1}' for i in range(4)]
    values = [random.randint(10, 100) for _ in categories]
    colors = plt.cm.viridis(np.linspace(0.2, 0.9, len(categories)))
    
    # Create figure
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(categories, values, color=colors)
    
    # Add value labels on top of bars
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                f'{height:.0f}', ha='center', va='bottom')
    
    ax.set_title(title)
    ax.set_ylabel('Values')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    # Save the figure
    plt.tight_layout()
    fig.savefig(output_path, dpi=100, bbox_inches='tight')
    plt.close()


def generate_document(
    output_path: str,
    output_dir: str = 'output',
    num_charts: int = 3,
    image_dir: str = 'chart_images'
) -> None:
    """Generate a Word document with random charts.
    
    Args:
        output_path: Filename for the output document.
        output_dir: Directory where the output document will be saved.
        num_charts: Number of charts to generate.
        image_dir: Directory to save temporary chart images.
        
    Raises:
        Exception: If there's an error during document generation.
    """
    try:
        # Ensure paths are Path objects
        output_dir = Path(output_dir)
        image_dir = Path(image_dir)
        
        # Create directories if they don't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        image_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize a new document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Random Charts Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph('Generated on: ' + 'July 27, 2025')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add some space
        
        # Generate random charts
        chart_paths = []
        for i in range(num_charts):
            # Alternate between pie and bar charts
            if i % 2 == 0:
                chart_type = 'pie'
                chart_path = image_dir / f'pie_chart_{i+1}.png'
                generate_pie_chart(str(chart_path), f'Random Data Distribution {i+1}')
            else:
                chart_type = 'bar'
                chart_path = image_dir / f'bar_chart_{i+1}.png'
                generate_bar_chart(str(chart_path), f'Quarterly Performance {i+1}')
            
            chart_paths.append((chart_path, chart_type))
        
        # Add charts to document
        for i, (chart_path, chart_type) in enumerate(chart_paths):
            # Add a heading for the chart
            doc.add_heading(f'Figure {i+1}: {chart_type.capitalize()} Chart', level=2)
            
            # Add the image
            doc.add_picture(str(chart_path), width=Inches(6))
            
            # Add description
            para = doc.add_paragraph(f'This is a randomly generated {chart_type} chart showing sample data.')
            para.paragraph_format.space_after = Pt(24)  # Add some space after each chart
            
            # Add a page break after each chart except the last one
            if i < len(chart_paths) - 1:
                doc.add_page_break()
        
        # Save the document
        output_file = output_dir / output_path
        doc.save(str(output_file))
        logger.info(f"Document successfully generated: {output_file}")
        
        # Clean up temporary image files
        for chart_path, _ in chart_paths:
            try:
                os.remove(chart_path)
            except Exception as e:
                logger.warning(f"Could not remove temporary file {chart_path}: {e}")
        
    except Exception as e:
        logger.error(f"Error generating document: {e}", exc_info=True)
        raise


def main() -> None:
    """Main function to demonstrate document generation with random charts."""
    try:
        generate_document(
            output_path="random_charts.docx",
            output_dir="output",
            num_charts=4,  # Generate 2 pie charts and 2 bar charts
            image_dir="temp_charts"
        )
        print("Document generated successfully!")
    except Exception as e:
        logger.critical("Failed to generate document", exc_info=True)
        raise


if __name__ == "__main__":
    main()