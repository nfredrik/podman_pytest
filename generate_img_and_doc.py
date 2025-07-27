"""
Module for generating Word documents with embedded images using docxtpl.
"""
"""
Module for generating Word documents with embedded images using python-docx-template.
"""
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Union

from docxtpl import DocxTemplate
from docx.shared import Mm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def add_image_to_paragraph(paragraph, image_path, width_mm=None, height_mm=None):
    """Add an image to a paragraph.
    
    Args:
        paragraph: The paragraph to add the image to.
        image_path: Path to the image file.
        width_mm: Optional width in millimeters.
        height_mm: Optional height in millimeters.
    """
    run = paragraph.add_run()
    
    # Add the image
    if width_mm and height_mm:
        run.add_picture(image_path, width=Mm(width_mm), height=Mm(height_mm))
    elif width_mm:
        run.add_picture(image_path, width=Mm(width_mm))
    elif height_mm:
        run.add_picture(image_path, height=Mm(height_mm))
    else:
        run.add_picture(image_path)


def generate_document(
    template_path: str,
    output_path: str,
    image_path: str,
    image_size: Optional[Dict[str, int]] = None,
    context: Optional[Dict[str, Any]] = None
) -> None:
    """Generate a Word document with an embedded image.
    
    Args:
        template_path: Path to the Word template file.
        output_path: Path where to save the generated document.
        image_path: Path to the image to embed.
        image_size: Optional dict with 'width' and/or 'height' for the image.
        context: Additional context variables for the template.
        
    Raises:
        FileNotFoundError: If template or image file doesn't exist.
        ValueError: If template processing fails.
    """
    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    
    # Set default image size if not provided (in millimeters)
    if image_size is None:
        image_size = {'width': 50, 'height': 30}  # Default to 50x30mm
    
    try:
        # Create output directory if it doesn't exist
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Load the template
        doc = DocxTemplate(template_path)
        
        # Create a new paragraph for the image
        doc.add_paragraph()  # Add a blank paragraph before the image
        
        # Add the image
        add_image_to_paragraph(
            doc.add_paragraph(),
            image_path=image_path,
            width_mm=image_size.get('width'),
            height_mm=image_size.get('height')
        )
        
        # Add any additional context
        if context:
            doc.render(context)
        
        # Save the document
        doc.save(output_path)
        
        logger.info(f"Document successfully generated: {output_path}")
        
    except Exception as e:
        logger.error(f"Error generating document: {e}")
        raise


def main() -> None:
    """Example usage of the document generation function."""
    try:
        generate_document(
            template_path='template.docx',
            output_path='output/generated_document.docx',
            image_path='image.jpg',
            image_size={'width': 300, 'height': 200},
            context={
                'title': 'Sample Document',
                'description': 'This document was generated automatically.'
            }
        )
    except Exception as e:
        logger.critical(f"Failed to generate document: {e}", exc_info=True)
        raise


if __name__ == "__main__":
    main()